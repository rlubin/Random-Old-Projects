from docx import Document
from docx.shared import Inches

class DocxBuilder:
  # def __init__(self, files, path):
  def __init__(self, _list, path):
    '''
    takes File objects (files) and creates a .docx file (path)
    '''
    document = Document()
    for line in _list:
      document.add_paragraph(line)
    # for file in files:
    #   for line in file.contents:
    #     document.add_paragraph(line)
    document.save(path)
